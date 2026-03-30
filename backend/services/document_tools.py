"""
Service C — Document Tools (LLM-callable functions).

These tools operate directly on the physical .docx file on disk,
locating paragraphs via their UUID bookmarks.
"""
import os
from docx import Document
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from backend.models import File
from backend.chromadb_client import delete_file_chunks
from backend.services.ingestion import (
    get_paragraph_bookmark,
    chunk_paragraphs,
    add_paragraph_bookmark,
)
from backend.chromadb_client import add_chunks
import uuid


def _find_paragraph_by_bookmark(doc: Document, bookmark_name: str):
    """
    Locate a paragraph in the document by its UUID bookmark name.
    Returns the paragraph object or None.
    """
    for paragraph in doc.paragraphs:
        for bookmark in paragraph._p.findall(qn("w:bookmarkStart")):
            name = bookmark.get(qn("w:name"))
            if name == bookmark_name:
                return paragraph
    return None


def _get_file_path(db: Session, file_id: str) -> str:
    """Look up the physical file path from MySQL."""
    db_file = db.query(File).filter(File.id == file_id).first()
    if not db_file:
        raise FileNotFoundError(f"File {file_id} not found in database.")
    return db_file.local_storage_path


def edit_docx_paragraph(
    db: Session,
    file_id: str,
    paragraph_uuid: str,
    new_text: str,
    user_id: str,
) -> dict:
    """
    Edit a specific paragraph in a .docx file identified by its UUID bookmark.

    Strategy for preserving formatting:
    - Keep the first run's formatting properties
    - Clear all runs
    - Add new text with the preserved formatting
    """
    file_path = _get_file_path(db, file_id)
    doc = Document(file_path)

    target = _find_paragraph_by_bookmark(doc, paragraph_uuid)
    if target is None:
        return {
            "success": False,
            "error": f"Paragraph with UUID {paragraph_uuid} not found in document.",
        }

    # Preserve the first run's formatting if available
    original_font_props = None
    if target.runs:
        first_run = target.runs[0]
        r_pr = first_run._r.find(qn("w:rPr"))
        if r_pr is not None:
            from copy import deepcopy
            original_font_props = deepcopy(r_pr)

    print(f"📖 [DocTool] Editing paragraph {paragraph_uuid}. Found: {target.text[:50]}...")

    # Clear all existing runs
    for run in target.runs:
        run._r.getparent().remove(run._r)

    # Add new text with preserved formatting
    new_run = target.add_run(new_text)
    if original_font_props is not None:
        new_run._r.insert(0, original_font_props)

    # Save the modified document
    doc.save(file_path)

    # Re-index this file's chunks in ChromaDB
    _reindex_file(db, file_id, user_id, file_path)

    return {
        "success": True,
        "message": f"Paragraph updated successfully.",
        "paragraph_uuid": paragraph_uuid,
    }


def delete_paragraph(
    db: Session,
    file_id: str,
    paragraph_uuid: str,
    user_id: str,
) -> dict:
    """Delete a specific paragraph from the .docx file."""
    file_path = _get_file_path(db, file_id)
    doc = Document(file_path)

    target = _find_paragraph_by_bookmark(doc, paragraph_uuid)
    if target is None:
        return {
            "success": False,
            "error": f"Paragraph with UUID {paragraph_uuid} not found.",
        }

    print(f"🗑️ [DocTool] Deleting paragraph {paragraph_uuid}.")
    
    # Remove the paragraph element from its parent
    parent = target._p.getparent()
    parent.remove(target._p)

    doc.save(file_path)
    _reindex_file(db, file_id, user_id, file_path)

    return {
        "success": True,
        "message": "Paragraph deleted successfully.",
    }


def append_paragraph(
    db: Session,
    file_id: str,
    after_paragraph_uuid: str,
    text: str,
    user_id: str,
) -> dict:
    """Append a new paragraph after the specified paragraph."""
    file_path = _get_file_path(db, file_id)
    doc = Document(file_path)

    target = _find_paragraph_by_bookmark(doc, after_paragraph_uuid)
    if target is None:
        return {
            "success": False,
            "error": f"Paragraph with UUID {after_paragraph_uuid} not found.",
        }

    print(f"➕ [DocTool] Appending paragraph after {after_paragraph_uuid}.")
    
    # Create new paragraph element after the target
    from docx.oxml import OxmlElement
    new_p = OxmlElement("w:p")
    new_paragraph_uuid = str(uuid.uuid4())

    # Insert after the target paragraph
    target._p.addnext(new_p)

    # Now get the paragraph object from the document
    # We need to find it in doc.paragraphs
    for p in doc.paragraphs:
        if p._p is new_p:
            p.add_run(text)
            add_paragraph_bookmark(p, new_paragraph_uuid)
            break

    doc.save(file_path)
    _reindex_file(db, file_id, user_id, file_path)

    return {
        "success": True,
        "message": "Paragraph appended successfully.",
        "new_paragraph_uuid": new_paragraph_uuid,
    }


def _reindex_file(db: Session, file_id: str, user_id: str, file_path: str) -> None:
    """Re-extract, chunk, and embed the entire file after an edit."""
    doc = Document(file_path)

    paragraphs_data = []
    for paragraph in doc.paragraphs:
        para_uuid = get_paragraph_bookmark(paragraph)
        if para_uuid:
            paragraphs_data.append({
                "uuid": para_uuid,
                "text": paragraph.text,
            })

    chunks = chunk_paragraphs(paragraphs_data)

    # Clear old chunks and re-insert
    delete_file_chunks(file_id)

    if not chunks:
        return

    chunk_ids = []
    chunk_documents = []
    chunk_metadatas = []

    for i, chunk in enumerate(chunks):
        chunk_id = f"{file_id}_chunk_{i}"
        chunk_ids.append(chunk_id)
        chunk_documents.append(chunk["text"])
        chunk_metadatas.append({
            "user_id": user_id,
            "file_id": file_id,
            "paragraph_uuids": ",".join(chunk["paragraph_uuids"]),
        })

    add_chunks(
        ids=chunk_ids,
        documents=chunk_documents,
        metadatas=chunk_metadatas,
    )
