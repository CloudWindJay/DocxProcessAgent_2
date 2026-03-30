"""
Service A — The Ingestion Pipeline.

Triggered on file upload:
1. Save the raw .docx to disk
2. Create the MySQL file record
3. Parse & tag every paragraph with a UUID bookmark
4. Chunk text with tiktoken
5. Embed chunks into ChromaDB
"""
import os
import uuid
import copy
from io import BytesIO

import tiktoken
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from backend.config import settings
from backend.models import File
from backend.chromadb_client import add_chunks, delete_file_chunks


# ── Bookmark Helpers ──────────────────────────────────────

def _make_bookmark_id() -> str:
    """Generate a positive integer ID for the XML w:id attribute."""
    return str(uuid.uuid4().int & 0x7FFFFFFF)


def add_paragraph_bookmark(paragraph, bookmark_name: str) -> None:
    """
    Inject an invisible XML bookmark into a paragraph.
    The bookmark_name is the UUID we use to relocate this paragraph later.
    """
    p_element = paragraph._p

    bid = _make_bookmark_id()

    # Create <w:bookmarkStart w:id="..." w:name="..." />
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), bookmark_name)

    # Create <w:bookmarkEnd w:id="..." />
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)

    # Insert at the very beginning of the paragraph XML
    p_element.insert(0, start)
    # Insert end right after start
    p_element.insert(1, end)


def get_paragraph_bookmark(paragraph) -> str | None:
    """
    Read the UUID bookmark name from a paragraph, if present.
    Returns None if no bookmark found.
    """
    for bookmark in paragraph._p.findall(qn("w:bookmarkStart")):
        name = bookmark.get(qn("w:name"))
        # Skip Word's built-in bookmarks
        if name and not name.startswith("_"):
            return name
    return None


# ── Tiktoken Chunking ────────────────────────────────────

def chunk_paragraphs(
    paragraphs: list[dict],
    max_tokens: int = 500,
    encoding_name: str = "cl100k_base",
) -> list[dict]:
    """
    Group paragraphs into chunks of <= max_tokens.

    Each paragraph dict: {"uuid": str, "text": str}
    Returns list of chunk dicts: {
        "text": str,           # concatenated text
        "paragraph_uuids": list[str],  # which paragraphs are in this chunk
    }
    """
    enc = tiktoken.get_encoding(encoding_name)
    chunks = []
    current_text = ""
    current_uuids = []
    current_tokens = 0

    for para in paragraphs:
        text = para["text"].strip()
        if not text:
            continue

        para_tokens = len(enc.encode(text))

        # If a single paragraph exceeds max_tokens, it becomes its own chunk
        if para_tokens > max_tokens:
            # Flush current buffer first
            if current_text:
                chunks.append({
                    "text": current_text.strip(),
                    "paragraph_uuids": list(current_uuids),
                })
                current_text = ""
                current_uuids = []
                current_tokens = 0

            chunks.append({
                "text": text,
                "paragraph_uuids": [para["uuid"]],
            })
            continue

        # If adding this paragraph would exceed the limit, flush first
        if current_tokens + para_tokens > max_tokens:
            chunks.append({
                "text": current_text.strip(),
                "paragraph_uuids": list(current_uuids),
            })
            current_text = ""
            current_uuids = []
            current_tokens = 0

        current_text += f"[ParaID: {para['uuid']}] {text}\n"
        current_uuids.append(para["uuid"])
        current_tokens += para_tokens

    # Flush remaining
    if current_text.strip():
        chunks.append({
            "text": current_text.strip(),
            "paragraph_uuids": list(current_uuids),
        })

    return chunks


# ── Main Ingestion Pipeline ──────────────────────────────

def run_ingestion(
    db: Session,
    user_id: str,
    file_id: str,
    filename: str,
    file_bytes: bytes,
) -> File:
    """
    Full ingestion pipeline:
    1. Save .docx to disk
    2. Create MySQL record
    3. Parse & tag paragraphs with UUID bookmarks
    4. Chunk with tiktoken
    5. Embed into ChromaDB
    """

    # ── 1. Save raw file to disk ──
    user_dir = os.path.join(settings.UPLOAD_DIR, user_id)
    os.makedirs(user_dir, exist_ok=True)
    file_path = os.path.join(user_dir, f"{file_id}.docx")

    with open(file_path, "wb") as f:
        f.write(file_bytes)

    # ── 2. Create MySQL record ──
    db_file = File(
        id=file_id,
        user_id=user_id,
        filename=filename,
        local_storage_path=file_path,
    )
    db.add(db_file)
    db.commit()
    db.refresh(db_file)

    # ── 3. Parse & Tag paragraphs ──
    doc = Document(file_path)
    paragraphs_data = []

    for paragraph in doc.paragraphs:
        # Check if already bookmarked (re-upload scenario)
        existing_uuid = get_paragraph_bookmark(paragraph)
        if existing_uuid:
            para_uuid = existing_uuid
        else:
            para_uuid = str(uuid.uuid4())
            add_paragraph_bookmark(paragraph, para_uuid)

        paragraphs_data.append({
            "uuid": para_uuid,
            "text": paragraph.text,
        })

    # Save the bookmarked version back
    doc.save(file_path)

    # ── 4. Chunk paragraphs ──
    chunks = chunk_paragraphs(paragraphs_data)

    if not chunks:
        return db_file

    # ── 5. Embed into ChromaDB ──
    # Clear any existing chunks for this file (idempotent)
    delete_file_chunks(file_id)

    chunk_ids = []
    chunk_documents = []
    chunk_metadatas = []

    for i, chunk in enumerate(chunks):
        chunk_id = f"{file_id}_chunk_{i}"
        chunk_ids.append(chunk_id)
        chunk_documents.append(chunk["text"])
        chunk_metadatas.append({
            "user_id": user_id,
            "file_id": file_id,
            "paragraph_uuids": ",".join(chunk["paragraph_uuids"]),
        })

    add_chunks(
        ids=chunk_ids,
        documents=chunk_documents,
        metadatas=chunk_metadatas,
    )

    return db_file
